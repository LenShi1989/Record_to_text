import ctypes
import os
import subprocess
import sys
import shutil
from pathlib import Path

from docx import Document
from faster_whisper import WhisperModel
from PyQt5.QtCore import QThread, pyqtSignal
from PyQt5.QtGui import QIcon
from PyQt5.QtWidgets import (
    QApplication,
    QFileDialog,
    QLabel,
    QProgressBar,
    QPushButton,
    QMessageBox,
    QTextEdit,
    QVBoxLayout,
    QWidget,
)


qt_plugins = os.path.join(
    sys.prefix, "Lib", "site-packages", "PyQt5", "Qt5", "plugins")
qt_platforms = os.path.join(qt_plugins, "platforms")

os.environ["QT_PLUGIN_PATH"] = qt_plugins
os.environ["QT_QPA_PLATFORM_PLUGIN_PATH"] = qt_platforms


def resolve_ffmpeg_path(filename):
    candidates = []

    if getattr(sys, "frozen", False):
        candidates.append(Path(sys._MEIPASS) / "ffmpeg" / filename)

    project_dir = Path(__file__).resolve().parent
    candidates.append(project_dir / "ffmpeg" / filename)
    candidates.append(Path("C:/ffmpeg") / filename)

    system_ffmpeg = shutil.which(filename)
    if system_ffmpeg:
        return system_ffmpeg

    for candidate in candidates:
        if candidate.exists():
            return str(candidate)

    return None


def resource_path(relative_path):
    if getattr(sys, "frozen", False):
        return str(Path(sys._MEIPASS) / relative_path)
    return str(Path(__file__).resolve().parent / relative_path)


FFMPEG_PATH = resolve_ffmpeg_path("ffmpeg.exe")


def get_ffmpeg_missing_message():
    return (
        "找不到 ffmpeg.exe。\n\n"
        "請確認以下其中一種方式已完成：\n"
        "1. 將 ffmpeg.exe 放到專案的 ffmpeg 資料夾\n"
        "2. 將 ffmpeg 安裝後加入系統 PATH\n"
        "3. 將 ffmpeg 安裝到 C:\\ffmpeg\\ffmpeg.exe"
    )


def convert_audio_to_wav(source_path, target_path):
    if not FFMPEG_PATH:
        raise FileNotFoundError(
            "找不到 ffmpeg.exe。請把 ffmpeg.exe 放到專案的 ffmpeg 資料夾，"
            "或安裝 ffmpeg 並加入 PATH。"
        )

    startupinfo = None
    creationflags = 0

    if os.name == "nt":
        startupinfo = subprocess.STARTUPINFO()
        startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
        creationflags = subprocess.CREATE_NO_WINDOW

    command = [
        FFMPEG_PATH,
        "-y",
        "-i",
        str(source_path),
        "-ac",
        "1",
        "-ar",
        "16000",
        str(target_path),
    ]

    subprocess.run(
        command,
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        startupinfo=startupinfo,
        creationflags=creationflags,
    )


myappid = "JQuan.com.tw"
ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)


class TranscribeThread(QThread):
    progress = pyqtSignal(int)
    finished = pyqtSignal(str, str)

    def __init__(self, file_path, model_name="medium"):
        super().__init__()
        self.file_path = file_path
        self.model_name = model_name

    def run(self):
        wav_file = None

        try:
            self.progress.emit(10)
            file_path = Path(self.file_path)
            wav_file = file_path.with_name(f"{file_path.stem}_transcribed.wav")

            convert_audio_to_wav(file_path, wav_file)

            self.progress.emit(30)

            model = WhisperModel(
                self.model_name,
                device="cpu",
                compute_type="int8",
            )

            segments, _info = model.transcribe(
                str(wav_file),
                language="zh",
                beam_size=5,
                vad_filter=False,
            )

            text = "".join(segment.text for segment in segments)

            self.progress.emit(70)

            sentences = [
                sentence.strip()
                for sentence in (
                    text.replace("。", "\n")
                    .replace("！", "\n")
                    .replace("？", "\n")
                    .replace("；", "\n")
                    .replace("，", "\n")
                ).split("\n")
                if sentence.strip()
            ]
            outline = "\n".join(
                f"{index + 1}. {sentence}" for index, sentence in enumerate(sentences))

            self.progress.emit(100)
            self.finished.emit(text, outline)

        except Exception as exc:
            self.finished.emit(f"(發生錯誤: {exc})", "")

        finally:
            if wav_file and wav_file.exists():
                try:
                    wav_file.unlink()
                except OSError:
                    pass


class SpeechToTextApp(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("音檔轉文字 by Len")
        self.setGeometry(300, 200, 600, 550)

        layout = QVBoxLayout()

        self.label = QLabel("請選擇音檔 (m4a / wav / mp3)")
        layout.addWidget(self.label)

        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        layout.addWidget(self.progress_bar)

        self.text_area = QTextEdit()
        self.text_area.setReadOnly(True)
        layout.addWidget(self.text_area)

        self.btn_open = QPushButton("選擇音檔")
        self.btn_open.clicked.connect(self.open_file)
        layout.addWidget(self.btn_open)

        self.btn_transcribe = QPushButton("開始轉錄")
        self.btn_transcribe.clicked.connect(self.transcribe_audio)
        self.btn_transcribe.setEnabled(False)
        layout.addWidget(self.btn_transcribe)

        self.btn_export = QPushButton("匯出 Word")
        self.btn_export.clicked.connect(self.export_word)
        self.btn_export.setEnabled(False)
        layout.addWidget(self.btn_export)

        self.setLayout(layout)

        self.audio_file = None
        self.transcript_text = ""
        self.outline_text = ""

        self.check_ffmpeg_on_startup()

    def check_ffmpeg_on_startup(self):
        if FFMPEG_PATH:
            return

        message = get_ffmpeg_missing_message()
        self.text_area.setText(message)
        self.btn_transcribe.setEnabled(False)
        QMessageBox.warning(self, "缺少 ffmpeg", message)

    def open_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "選擇音檔",
            "",
            "音訊檔案 (*.m4a *.wav *.mp3 *.aac *.mp4);;所有檔案 (*)",
        )
        if file_path:
            self.audio_file = file_path
            self.label.setText(f"已選擇：{Path(file_path).name}")
            self.btn_transcribe.setEnabled(True)

    def transcribe_audio(self):
        if not self.audio_file:
            self.text_area.setText("請先選擇音檔。")
            return

        self.text_area.setText("轉錄中，請稍候...")
        self.progress_bar.setValue(0)
        self.btn_transcribe.setEnabled(False)
        self.btn_export.setEnabled(False)

        self.thread = TranscribeThread(self.audio_file, "medium")
        self.thread.progress.connect(self.progress_bar.setValue)
        self.thread.finished.connect(self.on_transcription_finished)
        self.thread.start()

    def on_transcription_finished(self, text, outline):
        self.transcript_text = text
        self.outline_text = outline
        self.text_area.setText(
            f"==== 逐字稿 ====\n{text}\n\n==== 重點整理 ====\n{outline}"
        )
        self.btn_transcribe.setEnabled(True)
        if text and not text.startswith("(發生錯誤"):
            self.btn_export.setEnabled(True)

    def export_word(self):
        if not self.transcript_text:
            self.text_area.setText("沒有可匯出的轉錄內容。")
            return

        save_path, _ = QFileDialog.getSaveFileName(
            self,
            "儲存 Word 檔",
            "轉錄結果.docx",
            "Word 檔案 (*.docx)",
        )
        if not save_path:
            return

        doc = Document()
        doc.add_heading("音檔轉文字結果", level=1)

        doc.add_heading("逐字稿", level=2)
        doc.add_paragraph(self.transcript_text)

        doc.add_heading("重點整理", level=2)
        for line in self.outline_text.split("\n"):
            doc.add_paragraph(line)

        doc.save(save_path)
        self.text_area.append(f"\n已匯出 Word 檔：{save_path}")


if __name__ == "__main__":
    app = QApplication(sys.argv)
    icon_path = resource_path("Record.ico")
    app.setWindowIcon(QIcon(icon_path))
    window = SpeechToTextApp()
    window.setWindowIcon(QIcon(icon_path))
    window.show()
    sys.exit(app.exec_())
